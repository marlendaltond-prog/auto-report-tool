#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown文档转换为Word文档的工具
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START


def convert_md_to_docx(md_path, docx_path):
    """
    将Markdown文件转换为Word文件
    
    参数:
        md_path: Markdown文件路径
        docx_path: 输出Word文件路径
    """
    # 创建文档
    doc = Document()
    
    # 读取Markdown文件
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 处理行
    for line in lines:
        line = line.rstrip()  # 去除末尾换行符
        
        if not line:  # 空行
            doc.add_paragraph()
            continue
        
        # 处理标题
        if line.startswith('#'):
            level = line.count('#')
            title = line.lstrip('# ')
            
            if level == 1:
                # 一级标题
                doc.add_heading(title, level=0)
            elif level == 2:
                # 二级标题
                doc.add_heading(title, level=1)
            elif level == 3:
                # 三级标题
                doc.add_heading(title, level=2)
            elif level == 4:
                # 四级标题
                doc.add_heading(title, level=3)
            elif level == 5:
                # 五级标题
                doc.add_heading(title, level=4)
            elif level == 6:
                # 六级标题
                doc.add_heading(title, level=5)
        
        # 处理代码块
        elif line.startswith('```'):
            # 查找代码块结束位置
            code_lines = []
            next_lines = lines[lines.index(line)+1:]
            for code_line in next_lines:
                if code_line.strip().startswith('```'):
                    break
                code_lines.append(code_line.rstrip())
            
            # 添加代码段落
            if code_lines:
                code_text = '\n'.join(code_lines)
                paragraph = doc.add_paragraph(code_text)
                # 设置代码样式
                run = paragraph.runs[0]
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        
        # 处理无序列表
        elif line.startswith('- ') or line.startswith('* '):
            content = line[2:]
            paragraph = doc.add_paragraph(content)
            paragraph.paragraph_format.left_indent = Inches(0.5)
        
        # 处理表格
        elif '|' in line and '-' in line and lines.index(line) < len(lines)-1:
            # 检查是否为表格分隔线
            next_line = lines[lines.index(line)+1]
            if next_line.strip().startswith('|') and '---' in next_line:
                # 提取表格数据
                table_rows = []
                table_rows.append([cell.strip() for cell in line.strip('|').split('|')])
                
                # 跳过分隔线
                lines_iter = iter(lines[lines.index(line)+2:])
                for table_line in lines_iter:
                    if '|' in table_line:
                        table_rows.append([cell.strip() for cell in table_line.strip('|').split('|')])
                    else:
                        break
                
                # 创建表格
                if table_rows:
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    for i, row_data in enumerate(table_rows):
                        row = table.rows[i]
                        for j, cell_data in enumerate(row_data):
                            cell = row.cells[j]
                            cell.text = cell_data
                            # 设置表格样式
                            if i == 0:  # 表头
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in paragraph.runs:
                                        run.bold = True
        
        # 处理普通段落
        else:
            doc.add_paragraph(line)
    
    # 保存文档
    doc.save(docx_path)
    print(f"转换完成！Word文件已保存至: {docx_path}")


if __name__ == "__main__":
    # 输入Markdown文件路径
    md_file = "自动化报表工具使用说明书.md"
    # 输出Word文件路径
    docx_file = "自动化报表工具使用说明书.docx"
    
    # 转换
    convert_md_to_docx(md_file, docx_file)