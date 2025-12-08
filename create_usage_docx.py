#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
直接创建自动化报表工具的Word使用说明书
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def create_usage_document():
    """
    创建自动化报表工具的Word使用说明书
    """
    # 创建文档
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('自动化报表工具 - AutoReport Pro 使用说明书', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. 工具介绍
    doc.add_heading('1. 工具介绍', level=1)
    doc.add_paragraph('AutoReport Pro 是一款功能强大的自动化报表生成工具，支持多种数据源、多种输出格式和灵活的数据处理能力。')
    doc.add_paragraph('')
    
    features = doc.add_paragraph('主要功能特点：')
    features.runs[0].bold = True
    doc.add_paragraph('• 📊 支持多种数据源：Excel、CSV、SQL数据库、API')
    doc.add_paragraph('• 📄 支持多种输出格式：Excel、PDF、HTML、邮件')
    doc.add_paragraph('• 🔧 可配置的数据处理：过滤、计算、图表生成')
    doc.add_paragraph('• 🎨 可配置的报表样式和模板')
    doc.add_paragraph('• ⏰ 支持定时执行和邮件发送')
    doc.add_paragraph('• 📧 支持报表自动发送到指定邮箱')
    doc.add_paragraph('')
    
    # 2. 安装要求
    doc.add_heading('2. 安装要求', level=1)
    
    sys_req = doc.add_paragraph('系统要求：')
    sys_req.runs[0].bold = True
    doc.add_paragraph('• Windows/macOS/Linux')
    doc.add_paragraph('• Python 3.7 或更高版本')
    doc.add_paragraph('')
    
    dep_req = doc.add_paragraph('依赖包安装：')
    dep_req.runs[0].bold = True
    doc.add_paragraph('工具需要以下第三方依赖包：')
    code1 = doc.add_paragraph('pip install pandas openpyxl sqlalchemy jinja2 reportlab requests')
    code1.runs[0].font.name = 'Courier New'
    code1.runs[0].font.size = Pt(10)
    doc.add_paragraph('')
    doc.add_paragraph('或者使用提供的 requirements.txt 文件：')
    code2 = doc.add_paragraph('pip install -r requirements.txt')
    code2.runs[0].font.name = 'Courier New'
    code2.runs[0].font.size = Pt(10)
    doc.add_paragraph('')
    
    # 3. 快速开始
    doc.add_heading('3. 快速开始', level=1)
    
    example1 = doc.add_paragraph('1. 使用命令行参数生成报表：')
    example1.runs[0].bold = True
    code3 = doc.add_paragraph('python auto_report.py --data data.xlsx --output reports --format excel,pdf')
    code3.runs[0].font.name = 'Courier New'
    code3.runs[0].font.size = Pt(10)
    doc.add_paragraph('')
    
    example2 = doc.add_paragraph('2. 使用配置文件生成报表：')
    example2.runs[0].bold = True
    code4 = doc.add_paragraph('python auto_report.py --config report_config.json')
    code4.runs[0].font.name = 'Courier New'
    code4.runs[0].font.size = Pt(10)
    doc.add_paragraph('')
    
    # 4. 使用流程
    doc.add_heading('4. 使用流程', level=1)
    
    step1 = doc.add_paragraph('1. 准备数据源')
    step1.runs[0].bold = True
    doc.add_paragraph('   • 确保数据源文件（Excel/CSV）格式正确')
    doc.add_paragraph('   • 检查数据完整性和格式一致性')
    doc.add_paragraph('   • 如果使用SQL或API数据源，确保连接信息正确')
    doc.add_paragraph('')
    
    step2 = doc.add_paragraph('2. 配置报表参数')
    step2.runs[0].bold = True
    doc.add_paragraph('   • 选择输出格式（Excel/PDF/HTML/邮件）')
    doc.add_paragraph('   • 设置输出目录')
    doc.add_paragraph('   • 配置数据处理规则（可选）')
    doc.add_paragraph('   • 配置报表样式和模板（可选）')
    doc.add_paragraph('')
    
    step3 = doc.add_paragraph('3. 运行工具')
    step3.runs[0].bold = True
    doc.add_paragraph('   • 使用命令行参数直接运行')
    doc.add_paragraph('   • 或使用配置文件运行')
    doc.add_paragraph('   • 检查运行日志和错误提示')
    doc.add_paragraph('')
    
    step4 = doc.add_paragraph('4. 查看和使用报表')
    step4.runs[0].bold = True
    doc.add_paragraph('   • 在输出目录查看生成的报表文件')
    doc.add_paragraph('   • 如果配置了邮件发送，检查收件邮箱')
    doc.add_paragraph('   • 验证报表数据准确性')
    doc.add_paragraph('')
    
    step5 = doc.add_paragraph('5. 高级配置（可选）')
    step5.runs[0].bold = True
    doc.add_paragraph('   • 配置定时执行')
    doc.add_paragraph('   • 设置自定义数据处理逻辑')
    doc.add_paragraph('   • 使用自定义报表模板')
    doc.add_paragraph('')
    
    # 5. 命令行参数说明
    doc.add_heading('5. 命令行参数说明', level=1)
    
    # 创建表格
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    # 设置表头
    headers = ['参数', '说明', '示例']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加数据
    data = [
        ['--help', '显示基本帮助信息', 'python auto_report.py --help'],
        ['--help-all', '显示详细帮助信息', 'python auto_report.py --help-all'],
        ['--config', '配置文件路径', 'python auto_report.py --config config.json'],
        ['--data', '数据源路径', 'python auto_report.py --data data.xlsx'],
        ['--output', '输出目录', 'python auto_report.py --output reports'],
        ['--format', '输出格式，多个用逗号分隔', 'python auto_report.py --format excel,pdf'],
    ]
    
    for i, row in enumerate(data):
        for j, item in enumerate(row):
            cell = table.cell(i+1, j)
            cell.text = item
            if j == 2:  # 示例列使用等宽字体
                cell.paragraphs[0].runs[0].font.name = 'Courier New'
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph('')
    
    # 6. 实际使用示例
    doc.add_heading('6. 实际使用示例', level=1)
    
    example3 = doc.add_paragraph('示例：使用Excel数据源生成Excel和PDF格式报表')
    example3.runs[0].bold = True
    code5code6 = doc.add_paragraph('python auto_report.py --data "302594156_按序号_大学生对新能源汽车购买意向调查研究_254_246.xlsx" --output reports --format excel,pdf')
    code5.runs[0].font.name = 'Courier New'
    code5.runs[0].font.size = Pt(10)
    doc.add_paragraph('')
    
    explanation = doc.add_paragraph('参数说明：')
    explanation.runs[0].bold = True
    doc.add_paragraph('   • --data：指定数据源文件路径，这里使用了完整的文件名')
    doc.add_paragraph('   • --output：设置输出目录为当前目录下的reports文件夹')
    doc.add_paragraph('   • --format：指定输出格式为Excel和PDF，用逗号分隔')
    doc.add_paragraph('')
    
    # 7. 注意事项
    doc.add_heading('7. 注意事项', level=1)
    doc.add_paragraph('• 确保数据源文件路径正确，文件名包含空格时需要用引号括起来')
    doc.add_paragraph('• 输出目录如果不存在，工具会自动创建')
    doc.add_paragraph('• 确保有足够的磁盘空间存储生成的报表文件')
    doc.add_paragraph('• 对于大型数据集，可能需要较长时间生成报表')
    doc.add_paragraph('• 使用API数据源时，确保网络连接正常且有访问权限')
    doc.add_paragraph('')
    
    # 8. 常见问题
    doc.add_heading('8. 常见问题', level=1)
    
    q1 = doc.add_paragraph('Q1: 工具启动时提示缺少依赖包怎么办？')
    q1.runs[0].bold = True
    doc.add_paragraph('A: 请按照提示安装所需的依赖包：')
    code6 = doc.add_paragraph('pip install pandas openpyxl sqlalchemy jinja2 reportlab requests')
    code6.runs[0].font.name = 'Courier New'
    code6.runs[0].font.size = Pt(10)
    doc.add_paragraph('')
    
    q2 = doc.add_paragraph('Q2: 如何生成多个格式的报表？')
    q2.runs[0].bold = True
    doc.add_paragraph('A: 使用 --format 参数，多个格式用逗号分隔：')
    code7 = doc.add_paragraph('python auto_report.py --format excel,pdf,html')
    code7.runs[0].font.name = 'Courier New'
    code7.runs[0].font.size = Pt(10)
    doc.add_paragraph('')
    
    q3 = doc.add_paragraph('Q3: 数据源文件路径包含空格怎么办？')
    q3.runs[0].bold = True
    doc.add_paragraph('A: 使用引号将文件路径括起来：')
    code8 = doc.add_paragraph('python auto_report.py --data "大学生对新能源汽车购买意向调查研究.xlsx"')
    code8.runs[0].font.name = 'Courier New'
    code8.runs[0].font.size = Pt(10)
    doc.add_paragraph('')
    
    # 9. 技术支持
    doc.add_heading('9. 技术支持', level=1)
    doc.add_paragraph('如果您在使用过程中遇到问题或有功能建议，请通过以下方式联系：')
    doc.add_paragraph('• 邮箱：support@autoreport.com')
    doc.add_paragraph('• 官方文档：https://autoreport.example.com/docs')
    doc.add_paragraph('• GitHub：https://github.com/autoreport/pro')
    
    # 保存文档
    doc.save('自动化报表工具使用说明书.docx')
    print('Word使用说明书已生成：自动化报表工具使用说明书.docx')


if __name__ == "__main__":
    create_usage_document()